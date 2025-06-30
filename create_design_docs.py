from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

def create_design_doc(module_name, output_path):
    doc = Document()
    
    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 添加大标题（宋体小一）
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run(f'{module_name}模块概要设计文档')
    title_run.font.name = '宋体'
    title_run.font.size = Pt(36)  # 小一号字体
    title_run.bold = True
    
    # 添加空行
    for _ in range(10):
        doc.add_paragraph()
    
    # 添加版本表格（四号宋体）
    table = doc.add_table(rows=2, cols=4, style='Table Grid')
    table.autofit = False
    table.allow_autofit = False
    
    # 设置表格宽度
    table.columns[0].width = Inches(1.5)  # 日期列
    table.columns[1].width = Inches(1.5)  # 版本列
    table.columns[2].width = Inches(1.5)  # 作者列
    table.columns[3].width = Inches(2.5)  # 说明列
    
    # 填充表头（四号宋体）
    headers = ['日期', '版本', '作者', '说明']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        paragraph = header_cells[i].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(header)
        run.font.name = '宋体'
        run.font.size = Pt(14)  # 四号字体
        run.bold = True
    
    # 填充第一行数据
    data = ['2024-03-21', '1.0', '系统设计组', '初始版本']
    data_cells = table.rows[1].cells
    for i, data_item in enumerate(data):
        paragraph = data_cells[i].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(data_item)
        run.font.name = '宋体'
        run.font.size = Pt(14)  # 四号字体
    
    # 添加分页符
    doc.add_page_break()

    # 目录（宋体小二）
    h1 = doc.add_heading('目录', level=1)
    h1.style.font.name = '宋体'
    h1.style.font.size = Pt(22)
    doc.add_paragraph()
    doc.add_page_break()

    # 一、功能设计说明（宋体小二）
    h1 = doc.add_heading('一、功能设计说明', level=1)
    h1.style.font.name = '宋体'
    h1.style.font.size = Pt(22)

    # 1.1 模块概述（宋体三号）
    h2 = doc.add_heading('1.1 模块概述', level=2)
    h2.style.font.name = '宋体'
    h2.style.font.size = Pt(16)
    
    # 1.2 设计目标
    h2 = doc.add_heading('1.2 设计目标', level=2)
    h2.style.font.name = '宋体'
    h2.style.font.size = Pt(16)

    # 1.3 适用范围
    h2 = doc.add_heading('1.3 适用范围', level=2)
    h2.style.font.name = '宋体'
    h2.style.font.size = Pt(16)

    # 根据不同模块添加具体内容
    module_descriptions = {
        'log': '''日志模块是系统的核心基础设施，为整个系统提供统一的日志记录服务。该模块采用异步设计，支持多种日志级别，具备日志分类、过滤、轮转等功能，并提供灵活的配置选项。通过高效的缓冲机制和写入策略，确保日志记录对系统性能的影响最小化。

设计目标：
1. 高性能
   - 异步写入机制，避免阻塞业务线程
   - 内存缓冲区优化，减少磁盘I/O
   - 批量写入策略，提高写入效率
   - 文件句柄复用，减少系统资源消耗

2. 高可靠性
   - 数据持久化保证
   - 故障恢复机制
   - 磁盘空间管理
   - 日志完整性校验

3. 易用性
   - 简单清晰的API接口
   - 灵活的配置选项
   - 丰富的日志格式
   - 完善的文档支持

4. 可扩展性
   - 插件化架构设计
   - 自定义日志处理器
   - 多种输出目标支持
   - 第三方集成能力

适用范围：
1. 系统运行日志
   - 应用程序运行状态记录
   - 系统性能指标监控
   - 资源使用情况跟踪
   - 系统启动关闭过程记录

2. 业务操作日志
   - 用户行为追踪
   - 业务流程记录
   - 数据变更记录
   - 审计日志记录

3. 安全相关日志
   - 访问控制日志
   - 安全事件记录
   - 异常行为监控
   - 安全审计追踪

4. 调试诊断日志
   - 程序调试信息
   - 错误堆栈跟踪
   - 性能分析数据
   - 系统诊断信息

主要功能：
1. 日志级别管理
   1.1 日志级别定义
       - DEBUG：调试信息，用于开发和测试
       - INFO：一般信息，记录正常操作
       - WARN：警告信息，潜在问题提示
       - ERROR：错误信息，操作失败记录
       - FATAL：致命错误，系统崩溃记录

   1.2 级别控制策略
       - 全局级别控制
       - 模块级别控制
       - 动态级别调整
       - 条件级别过滤

   1.3 级别继承机制
       - 父子模块级别继承
       - 覆盖规则定义
       - 动态继承调整
       - 继承链管理

2. 异步处理机制
   2.1 异步写入设计
       - 生产者消费者模型
       - 多级缓冲队列
       - 批量处理优化
       - 背压机制实现

   2.2 性能优化策略
       - 内存池管理
       - 零拷贝技术
       - 写入合并优化
       - CPU亲和性调度

   2.3 可靠性保证
       - 数据持久化
       - 故障恢复
       - 数据一致性
       - 异常处理机制

3. 日志轮转
   3.1 轮转策略
       - 大小轮转：按文件大小切分
       - 时间轮转：按时间周期切分
       - 混合策略：同时支持大小和时间
       - 自定义策略：支持扩展实现

   3.2 归档管理
       - 自动压缩
       - 定期清理
       - 归档命名
       - 存储路径管理

   3.3 并发控制
       - 文件锁机制
       - 原子性操作
       - 并发写入控制
       - 读写分离设计

4. 格式化功能
   4.1 内置格式模板
       - 标准格式模板
       - JSON格式模板
       - XML格式模板
       - 自定义格式模板

   4.2 变量替换支持
       - 时间日期变量
       - 环境变量
       - 上下文变量
       - 自定义变量

   4.3 格式化扩展
       - 自定义格式器
       - 格式转换器
       - 编码处理
       - 国际化支持

5. 性能优化
   5.1 缓冲区管理
       - 多级缓存设计
       - 缓存大小动态调整
       - 缓存淘汰策略
       - 内存使用优化

   5.2 I/O优化
       - 异步I/O
       - 文件预读
       - 写入合并
       - 磁盘调度优化

   5.3 资源管理
       - 内存池
       - 线程池
       - 文件句柄池
       - 连接池

6. 监控和告警
   6.1 状态监控
       - 写入性能监控
       - 资源使用监控
       - 错误率监控
       - 延迟监控

   6.2 告警机制
       - 阈值告警
       - 趋势告警
       - 异常告警
       - 容量告警

   6.3 监控指标
       - 吞吐量统计
       - 响应时间统计
       - 错误率统计
       - 资源使用统计

7. 安全特性
   7.1 访问控制
       - 文件权限控制
       - 用户认证
       - 角色授权
       - 访问审计

   7.2 数据安全
       - 敏感信息脱敏
       - 数据加密
       - 完整性校验
       - 安全传输

   7.3 审计功能
       - 操作审计
       - 安全审计
       - 合规审计
       - 审计报告生成

8. 配置管理
   8.1 配置项
       - 基础配置
       - 性能配置
       - 安全配置
       - 扩展配置

   8.2 配置方式
       - 文件配置
       - 环境变量
       - 命令行参数
       - 远程配置

   8.3 配置热更新
       - 动态加载
       - 实时生效
       - 版本管理
       - 回滚机制

9. 扩展机制
   9.1 插件系统
       - 插件接口定义
       - 插件生命周期
       - 插件依赖管理
       - 插件版本控制

   9.2 自定义处理器
       - 输入处理器
       - 过滤处理器
       - 格式化处理器
       - 输出处理器

   9.3 集成能力
       - 第三方日志系统集成
       - 监控系统集成
       - 分析系统集成
       - 告警系统集成

10. 工具支持
    10.1 开发工具
        - 日志分析工具
        - 性能测试工具
        - 配置验证工具
        - 调试工具

    10.2 运维工具
        - 日志收集工具
        - 日志查询工具
        - 统计分析工具
        - 监控工具

    10.3 管理工具
        - 配置管理工具
        - 权限管理工具
        - 审计工具
        - 报告生成工具''',
        'plugins': '''插件模块提供了一个可扩展的插件系统框架，支持动态加载和管理插件。该模块实现了插件生命周期管理、依赖关系处理、版本控制等核心功能，并提供了标准的插件开发接口。

主要功能：
1. 插件生命周期管理
   - 插件的加载和卸载
   - 启动和停止控制
   - 状态监控和管理

2. 依赖关系处理
   - 插件间依赖检查
   - 版本兼容性验证
   - 循环依赖检测

3. 配置管理
   - 插件配置加载
   - 配置热更新
   - 配置验证

4. 版本控制
   - 版本兼容性检查
   - 版本升级管理
   - 回滚支持

5. 插件接口规范
   - 标准化的插件接口
   - 事件处理机制
   - 资源管理接口''',
        'rw_split': '''读写分离模块实现了数据库访问的智能分流机制，通过分析SQL语句特征，将读写请求分发到不同的数据库节点。该模块支持多种分发策略，包括轮询、权重、负载等算法。

主要功能：
1. SQL解析和路由
   - SQL语句解析
   - 读写操作识别
   - 路由规则执行

2. 负载均衡
   - 多种均衡策略
   - 动态权重调整
   - 负载监控

3. 健康检查
   - 节点状态监控
   - 自动故障转移
   - 节点恢复处理

4. 连接池管理
   - 连接池优化
   - 连接复用
   - 超时处理

5. 监控统计
   - 性能指标收集
   - 请求统计分析
   - 异常监控'''
    }
    
    p = doc.add_paragraph(module_descriptions.get(module_name.lower(), ''))

    # 二、系统设计（宋体小二）
    h1 = doc.add_heading('二、系统设计', level=1)
    h1.style.font.name = '宋体'
    h1.style.font.size = Pt(22)

    # 2.1 总体架构（宋体三号）
    h2 = doc.add_heading('2.1 总体架构', level=2)
    h2.style.font.name = '宋体'
    h2.style.font.size = Pt(16)

    # 2.2 模块设计
    h2 = doc.add_heading('2.2 模块设计', level=2)
    h2.style.font.name = '宋体'
    h2.style.font.size = Pt(16)

    # 2.3 接口设计
    h2 = doc.add_heading('2.3 接口设计', level=2)
    h2.style.font.name = '宋体'
    h2.style.font.size = Pt(16)

    # 2.4 数据结构设计
    h2 = doc.add_heading('2.4 数据结构设计', level=2)
    h2.style.font.name = '宋体'
    h2.style.font.size = Pt(16)

    # 添加架构描述
    architecture_descriptions = {
        'log': '''总体架构：
日志模块采用分层架构设计，主要包含以下几个层次：

1. 接口层
   1.1 API接口设计
       - 日志记录接口
         * void log(Level level, const char* format, ...)
         * void logf(Level level, const char* format, va_list args)
         * void log_binary(Level level, const void* data, size_t size)
       
       - 配置接口
         * void set_level(Level level)
         * void set_format(const char* format)
         * void set_output(const char* path)

       - 控制接口
         * void start()
         * void stop()
         * void flush()
         * void rotate()

   1.2 回调接口
       - 错误处理回调
       - 状态变更回调
       - 监控数据回调
       - 告警触发回调

   1.3 扩展接口
       - 自定义格式化器
       - 自定义输出处理器
       - 自定义过滤器
       - 自定义压缩器

2. 控制层
   2.1 日志管理器
       - 初始化和清理
       - 配置加载和更新
       - 资源分配和回收
       - 状态管理和监控

   2.2 过滤器
       - 级别过滤
       - 模块过滤
       - 内容过滤
       - 自定义过滤

   2.3 格式化器
       - 标准格式化
       - JSON格式化
       - XML格式化
       - 自定义格式化

3. 缓冲层
   3.1 内存缓冲区
       - 环形缓冲区
       - 多级缓存
       - 内存池管理
       - 溢出处理

   3.2 批处理器
       - 数据批量处理
       - 定时刷新
       - 条件触发
       - 强制刷新

   3.3 队列管理
       - 优先级队列
       - 延迟队列
       - 死信队列
       - 重试队列

4. 存储层
   4.1 文件管理器
       - 文件创建和删除
       - 文件打开和关闭
       - 文件读写操作
       - 文件锁定机制

   4.2 轮转管理器
       - 轮转策略实现
       - 文件命名管理
       - 归档处理
       - 清理处理

   4.3 存储优化
       - 异步写入
       - 批量写入
       - 压缩存储
       - 索引优化

模块设计：
1. LogManager（日志管理器）
   - 职责：总体协调各个组件的工作
   - 关键方法：
     * initialize(): 初始化日志系统
     * shutdown(): 关闭日志系统
     * configure(): 加载配置
     * updateConfig(): 更新配置
   - 属性：
     * 配置信息
     * 组件实例
     * 运行状态
     * 性能指标

2. LogFormatter（格式化器）
   - 职责：将日志内容格式化为指定格式
   - 关键方法：
     * format(): 格式化日志内容
     * parse(): 解析格式模板
     * addVariable(): 添加变量
     * validate(): 验证格式
   - 属性：
     * 格式模板
     * 变量映射
     * 缓存配置
     * 编码设置

3. LogBuffer（缓冲区）
   - 职责：提供高效的内存缓冲机制
   - 关键方法：
     * write(): 写入数据
     * flush(): 刷新数据
     * resize(): 调整大小
     * clear(): 清空缓冲区
   - 属性：
     * 缓冲区大小
     * 使用情况
     * 阈值设置
     * 状态标志

4. LogWriter（写入器）
   - 职责：处理日志的持久化存储
   - 关键方法：
     * write(): 写入日志
     * sync(): 同步数据
     * rotate(): 轮转文件
     * cleanup(): 清理文件
   - 属性：
     * 文件句柄
     * 写入统计
     * 错误计数
     * 性能指标

接口设计：
1. 日志记录接口
```cpp
class ILogger {
public:
    // 基本日志接口
    virtual void log(Level level, const char* format, ...) = 0;
    virtual void logf(Level level, const char* format, va_list args) = 0;
    virtual void log_binary(Level level, const void* data, size_t size) = 0;

    // 便捷接口
    virtual void debug(const char* format, ...) = 0;
    virtual void info(const char* format, ...) = 0;
    virtual void warn(const char* format, ...) = 0;
    virtual void error(const char* format, ...) = 0;
    virtual void fatal(const char* format, ...) = 0;

    // 控制接口
    virtual void set_level(Level level) = 0;
    virtual void set_format(const char* format) = 0;
    virtual void set_output(const char* path) = 0;
    
    // 管理接口
    virtual void start() = 0;
    virtual void stop() = 0;
    virtual void flush() = 0;
    virtual void rotate() = 0;
};
```

2. 格式化接口
```cpp
class IFormatter {
public:
    virtual std::string format(const LogRecord& record) = 0;
    virtual void set_pattern(const std::string& pattern) = 0;
    virtual void add_variable(const std::string& name, 
                            std::function<std::string()> provider) = 0;
    virtual bool validate_pattern(const std::string& pattern) = 0;
};
```

3. 输出接口
```cpp
class IWriter {
public:
    virtual bool write(const char* data, size_t size) = 0;
    virtual bool flush() = 0;
    virtual bool rotate() = 0;
    virtual void close() = 0;
    virtual bool is_open() const = 0;
    virtual std::string get_current_file() const = 0;
};
```

数据结构设计：
1. 日志记录结构
```cpp
struct LogRecord {
    Level level;            // 日志级别
    uint64_t timestamp;     // 时间戳
    std::string module;     // 模块名
    std::string message;    // 日志内容
    std::string file;       // 源文件
    int line;              // 行号
    std::string function;   // 函数名
    std::thread::id thread_id; // 线程ID
    std::map<std::string, std::string> context; // 上下文信息
};
```

2. 配置结构
```cpp
struct LogConfig {
    Level level;           // 日志级别
    std::string format;    // 格式模板
    std::string output;    // 输出路径
    size_t buffer_size;    // 缓冲区大小
    size_t flush_interval; // 刷新间隔
    size_t max_file_size;  // 最大文件大小
    int max_files;         // 最大文件数
    bool async;            // 是否异步
    bool compress;         // 是否压缩
    std::map<std::string, std::string> extra; // 扩展配置
};
```

3. 统计信息结构
```cpp
struct LogStats {
    uint64_t total_records;    // 总记录数
    uint64_t total_bytes;      // 总字节数
    uint64_t flush_count;      // 刷新次数
    uint64_t rotate_count;     // 轮转次数
    uint64_t error_count;      // 错误次数
    uint64_t dropped_count;    // 丢弃次数
    double avg_latency;        // 平均延迟
    std::map<Level, uint64_t> level_stats; // 级别统计
};
```

性能设计：
1. 写入性能
   - 目标：单线程写入速度 > 100,000条/秒
   - 批量写入提升：50%
   - 延迟要求：P99 < 1ms
   - 内存占用：< 100MB

2. 资源消耗
   - CPU使用率：< 5%
   - 内存使用率：< 1%
   - 磁盘I/O：< 10%
   - 网络带宽：< 5%

3. 可靠性指标
   - 可用性：99.999%
   - 数据丢失率：< 0.001%
   - 错误率：< 0.0001%
   - MTBF：> 10000小时

4. 扩展性指标
   - 单文件最大支持：100GB
   - 最大并发线程：1000
   - 配置项数量：>100
   - 插件支持：>20种

安全设计：
1. 访问控制
   - 文件权限：644
   - 目录权限：755
   - 用户权限：最小权限
   - 进程权限：非root

2. 数据安全
   - 敏感信息加密
   - 传输加密：TLS 1.3
   - 存储加密：AES-256
   - 完整性校验：SHA-256

3. 审计功能
   - 操作审计
   - 访问审计
   - 安全审计
   - 合规审计

异常处理：
1. 系统异常
   - 磁盘满：自动清理
   - 内存不足：降级服务
   - CPU高负载：任务队列
   - 网络中断：本地缓存

2. 业务异常
   - 格式错误：默认格式
   - 配置错误：默认配置
   - 参数错误：参数校验
   - 状态错误：状态恢复

3. 恢复机制
   - 自动重试
   - 备份恢复
   - 降级服务
   - 熔断保护

测试设计：
1. 单元测试
   - 接口测试
   - 功能测试
   - 性能测试
   - 压力测试

2. 集成测试
   - 模块集成
   - 系统集成
   - 性能集成
   - 压力集成

3. 验收测试
   - 功能验收
   - 性能验收
   - 安全验收
   - 可靠性验收''',
        'plugins': '''插件模块采用微内核架构设计，主要包含以下部分：

1. 核心层
   - 插件生命周期管理
   - 提供基础服务接口
   - 处理插件通信

2. 管理层
   - 依赖关系管理
   - 版本控制
   - 配置管理

3. 接口层
   - 标准插件接口
   - 扩展点定义
   - 事件处理

4. 资源层
   - 资源隔离
   - 权限控制
   - 资源管理''',
        'rw_split': '''读写分离模块采用分层架构设计，主要包含以下层次：

1. 接入层
   - SQL请求接收
   - 协议解析
   - 连接管理

2. 解析层
   - SQL解析
   - 语义分析
   - 操作识别

3. 路由层
   - 路由规则处理
   - 负载均衡
   - 节点选择

4. 执行层
   - 请求转发
   - 结果处理
   - 异常处理'''
    }
    
    p = doc.add_paragraph(architecture_descriptions.get(module_name.lower(), ''))

    # 三、接口设计（宋体小二）
    h1 = doc.add_heading('三、接口设计', level=1)
    h1.style.font.name = '宋体'
    h1.style.font.size = Pt(22)

    # 四、性能设计（宋体小二）
    h1 = doc.add_heading('四、性能设计', level=1)
    h1.style.font.name = '宋体'
    h1.style.font.size = Pt(22)

    # 五、安全设计（宋体小二）
    h1 = doc.add_heading('五、安全设计', level=1)
    h1.style.font.name = '宋体'
    h1.style.font.size = Pt(22)

    # 六、测试设计（宋体小二）
    h1 = doc.add_heading('六、测试设计', level=1)
    h1.style.font.name = '宋体'
    h1.style.font.size = Pt(22)

    # 保存文档
    doc.save(output_path)
    print(f"已生成{module_name}模块设计文档")

def main():
    # 模块名称和中文映射
    module_names = {
        'log': '日志',
        'plugins': '插件',
        'rw_split': '读写分离',
        'security': '安全',
        'shard': '分片',
        'sql_extend': 'SQL扩展',
        'storage': '存储',
        'transaction': '事务'
    }
    
    # 生成所有模块的设计文档
    for module, module_cn in module_names.items():
        output_path = f'nuxx_code/{module}/{module_cn}模块概要设计文档.docx'
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        create_design_doc(module, output_path)
        print(f"已生成{module_cn}模块设计文档")

if __name__ == '__main__':
    main() 