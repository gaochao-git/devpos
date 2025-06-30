from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_formatted_doc():
    doc = Document()
    
    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 添加大标题
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.add_run('日志模块概要设计文档')
    title_run.font.name = '黑体'
    title_run.font.size = Pt(26)
    title_run.bold = True
    
    # 添加空行
    for _ in range(10):
        doc.add_paragraph()
    
    # 添加版本表格
    table = doc.add_table(rows=2, cols=4, style='Table Grid')
    table.autofit = False
    table.allow_autofit = False
    
    # 设置表格宽度
    table.columns[0].width = Inches(1.5)  # 日期列
    table.columns[1].width = Inches(1.5)  # 版本列
    table.columns[2].width = Inches(1.5)  # 作者列
    table.columns[3].width = Inches(2.5)  # 说明列
    
    # 填充表头
    headers = ['日期', '版本', '作者', '说明']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        paragraph = header_cells[i].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(header)
        run.font.name = '宋体'
        run.font.size = Pt(12)
        run.bold = True
    
    # 填充第一行数据
    data = ['2024-03-21', '1.0', '系统设计组', '']
    data_cells = table.rows[1].cells
    for i, data_item in enumerate(data):
        paragraph = data_cells[i].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run(data_item)
        run.font.name = '宋体'
        run.font.size = Pt(12)
    
    # 添加分页符
    doc.add_page_break()

    # 添加目录标题
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    toc_run = toc_heading.add_run('目录')
    toc_run.font.name = '黑体'
    toc_run.font.size = Pt(16)
    toc_run.bold = True
    
    # 添加目录（这里只是示意，实际目录会自动生成）
    doc.add_paragraph('1. 引言.....................................................1')
    doc.add_paragraph('2. 总体设计.................................................2')
    doc.add_paragraph('3. 详细设计.................................................3')
    doc.add_paragraph('4. 接口设计.................................................4')
    doc.add_paragraph('5. 性能设计.................................................5')
    
    # 添加分页符
    doc.add_page_break()

    # 1. 引言
    h1 = doc.add_heading('1. 引言', level=1)
    h1.style.font.name = '黑体'
    h1.style.font.size = Pt(16)
    
    h2 = doc.add_heading('1.1 编写目的', level=2)
    h2.style.font.name = '黑体'
    h2.style.font.size = Pt(14)
    p = doc.add_paragraph('本文档旨在详细描述日志模块的设计方案，为开发人员提供实现指导，为测试人员提供测试依据。')
    p.style.font.name = '宋体'
    p.style.font.size = Pt(12)
    
    h2 = doc.add_heading('1.2 背景', level=2)
    p = doc.add_paragraph('日志模块是系统中的核心基础模块，为所有其他模块提供日志记录服务。该模块需要保证高性能、高可靠性，并支持异步操作。')
    
    # 2. 总体设计
    h1 = doc.add_heading('2. 总体设计', level=1)
    h2 = doc.add_heading('2.1 设计目标', level=2)
    p = doc.add_paragraph('主要目标包括：')
    doc.add_paragraph('• 提供异步日志记录功能', style='List Bullet')
    doc.add_paragraph('• 支持多种日志级别', style='List Bullet')
    doc.add_paragraph('• 确保线程安全', style='List Bullet')
    doc.add_paragraph('• 支持日志文件轮转', style='List Bullet')
    doc.add_paragraph('• 提供灵活的配置选项', style='List Bullet')
    
    h2 = doc.add_heading('2.2 系统架构', level=2)
    p = doc.add_paragraph('日志模块采用分层架构设计：')
    doc.add_paragraph('• 接口层：提供日志记录API', style='List Bullet')
    doc.add_paragraph('• 控制层：处理日志格式化和路由', style='List Bullet')
    doc.add_paragraph('• 存储层：负责日志持久化', style='List Bullet')
    
    # 3. 详细设计
    h1 = doc.add_heading('3. 详细设计', level=1)
    h2 = doc.add_heading('3.1 核心类设计', level=2)
    
    # 添加类图表格
    table = doc.add_table(rows=4, cols=2, style='Table Grid')
    table.autofit = False
    
    # 设置表格标题
    cells = table.rows[0].cells
    cells[0].text = '类名'
    cells[1].text = '主要职责'
    
    # 填充类信息
    data = [
        ['LogManager', '日志管理器，负责初始化和配置管理'],
        ['LogWriter', '日志写入器，处理异步写入操作'],
        ['LogFormatter', '日志格式化器，处理日志格式转换']
    ]
    
    for i, row in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = row[0]
        cells[1].text = row[1]
    
    h2 = doc.add_heading('3.2 关键流程', level=2)
    p = doc.add_paragraph('日志记录的主要流程包括：')
    doc.add_paragraph('1. 应用程序调用日志接口', style='List Number')
    doc.add_paragraph('2. LogManager接收并验证日志信息', style='List Number')
    doc.add_paragraph('3. LogFormatter进行格式化处理', style='List Number')
    doc.add_paragraph('4. LogWriter异步写入日志文件', style='List Number')
    
    # 4. 接口设计
    h1 = doc.add_heading('4. 接口设计', level=1)
    h2 = doc.add_heading('4.1 外部接口', level=2)
    
    # 添加接口表格
    table = doc.add_table(rows=4, cols=3, style='Table Grid')
    table.autofit = False
    
    # 设置表格标题
    cells = table.rows[0].cells
    cells[0].text = '接口名称'
    cells[1].text = '参数'
    cells[2].text = '说明'
    
    # 填充接口信息
    data = [
        ['log()', 'level, message', '记录指定级别的日志'],
        ['setLogLevel()', 'level', '设置日志级别'],
        ['setLogPath()', 'path', '设置日志文件路径']
    ]
    
    for i, row in enumerate(data, 1):
        cells = table.rows[i].cells
        cells[0].text = row[0]
        cells[1].text = row[1]
        cells[2].text = row[2]
    
    # 5. 性能设计
    h1 = doc.add_heading('5. 性能设计', level=1)
    h2 = doc.add_heading('5.1 性能指标', level=2)
    p = doc.add_paragraph('性能目标：')
    doc.add_paragraph('• 日志写入延迟：<100ms', style='List Bullet')
    doc.add_paragraph('• 每秒处理能力：>10000条日志', style='List Bullet')
    doc.add_paragraph('• 内存占用：<50MB', style='List Bullet')
    
    h2 = doc.add_heading('5.2 优化措施', level=2)
    p = doc.add_paragraph('为达到性能目标，采取以下优化措施：')
    doc.add_paragraph('• 使用无锁队列', style='List Bullet')
    doc.add_paragraph('• 实现批量写入', style='List Bullet')
    doc.add_paragraph('• 采用内存映射文件', style='List Bullet')
    doc.add_paragraph('• 优化字符串处理', style='List Bullet')

    # 设置所有段落的行距为1.5倍
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
        # 设置首行缩进（除了标题）
        if not paragraph.style.name.startswith('Heading'):
            paragraph.paragraph_format.first_line_indent = Pt(24)
    
    # 保存文档
    doc.save('nuxx_code/log/日志模块概要设计文档.docx')

if __name__ == '__main__':
    create_formatted_doc() 